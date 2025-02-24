from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font
import os

# Function to create a new Word document with title and subtitle
# 创建一个带有标题和副标题的新Word文档的函数
def create_word_document(file_path, title, subtitle, content, 
                         title_font='Arial', title_size=24, title_color=(0, 0, 0),
                         subtitle_font='Arial', subtitle_size=18, subtitle_color=(0, 0, 0),
                         content_font='Arial', content_size=12, content_color=(0, 0, 0)):
    try:
        doc = Document()
        # print(f"当前目录地址: {os.getcwd()}")
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title_paragraph.runs:
            title_run = title_paragraph.runs[0]
            title_run.font.name = title_font
            title_run.font.size = Pt(title_size)
            title_run.font.color.rgb = RGBColor(*title_color)

        subtitle_paragraph = doc.add_heading(subtitle, level=2)
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if subtitle_paragraph.runs:
            subtitle_run = subtitle_paragraph.runs[0]
            subtitle_run.font.name = subtitle_font
            subtitle_run.font.size = Pt(subtitle_size)
            subtitle_run.font.color.rgb = RGBColor(*subtitle_color)

        content_paragraph = doc.add_paragraph(content)
        if content_paragraph.runs:
            content_run = content_paragraph.runs[0]
            content_run.font.name = content_font
            content_run.font.size = Pt(content_size)
            content_run.font.color.rgb = RGBColor(*content_color)

        doc.save(file_path)
    except Exception as e:
        print(f"创建Word文档时发生错误: {e}")
        print(f"参数: file_path={file_path}, title={title}, subtitle={subtitle}, content={content}")
        raise

# Function to modify an existing Word document with additional content
# 修改现有Word文档并添加内容的函数
def modify_word_document(file_path, content, 
                         content_font='Arial', content_size=12, content_color=(0, 0, 0)):
    try:
        doc = Document(file_path)
        content_paragraph = doc.add_paragraph(content)
        content_run = content_paragraph.runs[0]
        content_run.font.name = content_font
        content_run.font.size = Pt(content_size)
        content_run.font.color.rgb = RGBColor(*content_color)
        doc.save(file_path)
    except Exception as e:
        print(f"修改Word文档时发生错误: {e}")
        print(f"参数: file_path={file_path}, content={content}")
        raise

# Function to create a new Excel workbook with title and data
# 创建一个带有标题和数据的新Excel工作簿的函数
def create_excel_workbook(file_path, sheet_name, title, data, 
                          title_font='Arial', title_size=14, title_color='000000'):
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = Font(name=title_font, size=title_size, bold=True, color=title_color)

        for row in data:
            ws.append(row)

        wb.save(file_path)
    except Exception as e:
        print(f"创建Excel工作簿时发生错误: {e}")
        print(f"参数: file_path={file_path}, sheet_name={sheet_name}, title={title}, data={data}")
        raise

# Function to modify an existing Excel workbook with additional data
# 修改现有Excel工作簿并添加数据的函数
def modify_excel_workbook(file_path, sheet_name, data):
    try:
        wb = load_workbook(file_path)
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
        else:
            ws = wb.create_sheet(sheet_name)

        for row in data:
            ws.append(row)

        wb.save(file_path)
    except Exception as e:
        print(f"修改Excel工作簿时发生错误: {e}")
        print(f"参数: file_path={file_path}, sheet_name={sheet_name}, data={data}")
        raise

# General function to handle both Word and Excel documents
# 通用函数来处理Word和Excel文档
def handle_document(doc_type, action, file_path, **kwargs):
    try:
        if doc_type == 'word':
            if action == 'create':
                create_word_document(file_path, **kwargs)
            elif action == 'modify':
                modify_word_document(file_path, **kwargs)
            else:
                raise ValueError(f"Unsupported action '{action}' for document type 'word'")
        elif doc_type == 'excel':
            if action == 'create':
                create_excel_workbook(file_path, **kwargs)
            elif action == 'modify':
                modify_excel_workbook(file_path, **kwargs)
            else:
                raise ValueError(f"Unsupported action '{action}' for document type 'excel'")
        else:
            raise ValueError(f"Unsupported document type '{doc_type}'")
    except Exception as e:
        print(f"An error occurred: {e}")
        raise

# Example usage of the functions
# 函数的示例用法

# # 示例：创建一个新的Word文档
# create_word_document(
#     file_path='example.docx',
#     title='主标题',
#     subtitle='副标题',
#     content='这是文档的内容。'
# )

# # 示例：修改现有的Word文档
# modify_word_document(
#     file_path='example.docx',
#     content='这是添加到文档的新内容。'
# )

# # 示例：创建一个新的Excel工作簿
# create_excel_workbook(
#     file_path='example.xlsx',
#     sheet_name='Sheet1',
#     title='工作簿标题',
#     data=[
#         ['列1', '列2', '列3'],
#         [1, 2, 3],
#         [4, 5, 6]
#     ]
# )

# # 示例：修改现有的Excel工作簿
# modify_excel_workbook(
#     file_path='example.xlsx',
#     sheet_name='Sheet1',
#     data=[
#         [7, 8, 9],
#         [10, 11, 12]
#     ]
# )
# doc_type = 'word'
# action = 'create'
# file_path = 'data/plugins/astrbot_plugin_miaomiao/gen_doc/ai_understanding.docx'
# file_path = file_path.replace("data/plugins/astrbot_plugin_miaomiao/", "")
# kwargs = {'file_path': file_path,  'title': 'AI 数字比较的误判', 'content': 'AI 在处理数字时，如果数据类型不正确或存在精度问题，可能会导致错误的比较结果。例如，如果数字被当作字符串处理，那么字符串的比较是从左到右逐位比较的，因此 "7.11" 会被认为小于 "7.8"，因为第一个不同的字符是 "1" 和 "8"，而 "1" 小于 "8"。此外，浮点数在计算机中的存储方式也可能导致精度损失，从而影响比较结果。', 'subtitle': '为什么 AI 认为 7.11 小于 7.8'}
# handle_document(doc_type, action, **kwargs)